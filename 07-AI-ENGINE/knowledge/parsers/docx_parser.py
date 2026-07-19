"""
app/knowledge/parsers/docx_parser.py
────────────────────────────────────────────────────────────────
Microsoft Word (.docx) document parser.

Extracts:
  - Body paragraphs (in order)
  - Table cell content
  - Heading text
  - Bold/italic text (as plain text)

Skips: images, embedded objects, comments, tracked changes.

Requires: python-docx
"""

from __future__ import annotations

import io
import logging

from knowledge.parsers.base import BaseParser, ParseResult

log = logging.getLogger(__name__)


class DocxParser(BaseParser):
    supported_extensions = [".docx"]
    doc_type = "docx"

    def parse(self, file_bytes: bytes, filename: str = "") -> ParseResult:
        try:
            import docx
        except ImportError:
            return ParseResult(
                success=False,
                error="python-docx not installed. Run: pip install python-docx",
            )

        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            sections: list[str] = []

            # ── Extract body paragraphs ───────────────
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Mark headings
                if para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    prefix = "#" * min(int(level), 6) if level.isdigit() else "##"
                    sections.append(f"{prefix} {text}")
                else:
                    sections.append(text)

            # ── Extract tables ────────────────────────
            for table in doc.tables:
                rows: list[str] = []
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        rows.append(" | ".join(cells))
                if rows:
                    sections.append("\n".join(rows))

            content = "\n\n".join(sections)

            if not content.strip():
                return ParseResult(
                    success=False,
                    error="DOCX file contains no extractable text",
                )

            # Use filename as title if available
            title = filename.replace(".docx", "").replace("_", " ").strip()
            if not title:
                title = "Word Document"

            return ParseResult(
                success=True,
                content=content,
                title=title,
                doc_type=self.doc_type,
                page_count=len(doc.sections),
                word_count=len(content.split()),
                extra_info={
                    "paragraph_count": len(doc.paragraphs),
                    "table_count":     len(doc.tables),
                },
            )

        except Exception as exc:
            log.error("DOCX parse failed: %s", exc)
            return ParseResult(
                success=False,
                error=f"DOCX parse error: {exc}",
            )


docx_parser = DocxParser()
